# import fitz 
# import docx
# import re

# '''
# Method to extract text from PDF
# @param pdf_path: path to the PDF file
# @return: text extracted from the PDF
# '''
# def extract_text_from_pdf(pdf_path):
#     text=""
#     doc=fitz.open(pdf_path)
#     for page in doc:
#         text+=page.get_text("text")
#     return text

# '''
# Method to extract text from DOCX
# @param docx_path: path to the DOCX file
# @return: text extracted from the DOCX
# '''
# def extract_text_from_docx(docx_path):
#     doc=docx.Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# '''
# Method to extract data from the CV
# @param file_path: path to the CV file
# @return: data extracted from the CV
# '''
# def extract_cv_data(file_path):
#     if file_path.endswith(".pdf"):
#         text=extract_text_from_pdf(file_path)
#     elif file_path.endswith(".docx"):
#         text=extract_text_from_docx(file_path)
#     else:
#         return {}

#     #extracting data
#     # cv_data = {
#     #     "personal_info":{"name": "John Doe", "email": "johndoe@example.com"},
#     #     "education":["Bachelor in Computer Science"],
#     #     "certifications":["Python", "Machine Learning"],
#     #     "projects":["AI Resume Parser"]
#     # }

#     #reegex to extract Name, mail
#     name_match=re.search(r"([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)", text)
#     email_match=re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)

#     #name
#     name= name_match.group(0) if name_match else "Unknown"
#     #email
#     email =email_match.group(0) if email_match else "Unknown"




#     # #Education
#     # education_section=re.findall(r"(Bachelor|Master|Ph\.?D|B\.?Sc|BSc|M\.?Sc|Diploma|Associate) in [A-Za-z\s]+", text)
#     # if not education_section:
#     #     # Try alternative format:"Degree: [degree]"
#     #     education_section=re.findall(r"Degree:\s*(.*)", text)

#     # #Certifications
#     # certifications_section=re.findall(r"(Certified|Certificate|Diploma) in [A-Za-z\s]+", text)
#     # if not certifications_section:
        
#     #     certs_match=re.search(r"Certifications:\s*(.*)", text)
#     #     if certs_match:
#     #         certifications_section=certs_match.group(1).split(",")

#     # #projects
#     # project_lines=[line.strip() for line in text.split("\n") if "project" in line.lower()]

   
#     # cv_data={
#     #     "personal_info":{"name":name,"email":email},
#     #     "education":education_section,
#     #     "certifications":certifications_section,
#     #     "projects":project_lines
#     # }

#     # return cv_data

#     sections = {
#         "Education": [],
#         "Certifications": [],
#         "Projects": []
#     }

#     current_section = None
#     lines = text.split("\n")

#     for line in lines:
#         line = line.strip()

#         if re.search(r"\b(Education|Certifications|Projects|Work Experience|Skills)\b", line, re.IGNORECASE):
#             current_section = line.lower()

#         elif current_section:
#             if current_section.startswith("education"):
#                 sections["Education"].append(line)
#             elif current_section.startswith("certifications"):
#                 sections["Certifications"].append(line)
#             elif current_section.startswith("projects"):
#                 sections["Projects"].append(line)

#     # **Clean extracted data**
#     education = [edu for edu in sections["Education"] if len(edu) > 3]
#     certifications = [cert for cert in sections["Certifications"] if len(cert) > 3]
#     projects = [proj for proj in sections["Projects"] if len(proj) > 3]

#     # **Return Extracted Data**
#     cv_data = {
#         "personal_info": {"name": name, "email": email},
#         "education": education,
#         "certifications": certifications,
#         "projects": projects
#     }
    

#     return cv_data

import fitz  # PyMuPDF for PDFs
import docx  # For DOCX files
import re

'''
Method to extract text from PDF
@param pdf_path: path to the PDF file
@return: extracted text
'''
def extract_text_from_pdf(pdf_path):
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        text += page.get_text("text") + "\n"
    return text.strip()

'''
Method to extract text from DOCX
@param docx_path: path to the DOCX file
@return: extracted text
'''
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()

'''
Method to extract data from the CV
@param file_path: path to the CV file
@return: structured data extracted from the CV
'''
def extract_cv_data(file_path):
    if file_path.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        return {}

    # Extracting Name (First Capitalized Words at the Start)
    name_match = re.search(r"(?m)^\s*([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)", text)
    name = name_match.group(0) if name_match else "Unknown"

    # Extracting Email
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    email = email_match.group(0) if email_match else "Unknown"

    # Extracting Phone Number
    phone_match = re.search(r"\+?\d[\d\s\-\(\)]{8,}\d", text)
    phone = phone_match.group(0) if phone_match else "Unknown"

    # Extracting LinkedIn & GitHub
    linkedin_match = re.search(r"(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+", text)
    github_match = re.search(r"(https?://)?(www\.)?github\.com/[a-zA-Z0-9_-]+", text)
    
    linkedin = linkedin_match.group(0) if linkedin_match else "Unknown"
    github = github_match.group(0) if github_match else "Unknown"

    # Extract Sections (Education, Certifications, Projects)
    sections = {"Education": [], "Certifications": [], "Projects": []}
    current_section = None

    lines = text.split("\n")
    for line in lines:
        line = line.strip()

        # Check for section headers
        if re.search(r"\b(Education|Certifications|Projects|Work Experience|Skills)\b", line, re.IGNORECASE):
            current_section = line.lower()

        elif current_section:
            if "education" in current_section:
                sections["Education"].append(line)
            elif "certifications" in current_section:
                sections["Certifications"].append(line)
            elif "projects" in current_section:
                sections["Projects"].append(line)

    # Clean extracted data
    education = [edu for edu in sections["Education"] if len(edu) > 3]
    certifications = [cert for cert in sections["Certifications"] if len(cert) > 3]
    projects = [proj for proj in sections["Projects"] if len(proj) > 3]

    # Return extracted CV data
    cv_data = {
        "personal_info": {
            "name": name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin,
            "github": github
        },
        "education": education,
        "certifications": certifications,
        "projects": projects
    }

    return cv_data
